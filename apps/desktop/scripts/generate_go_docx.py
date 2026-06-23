#!/usr/bin/env python3
"""Generate a ~400-word Go Language Basics DOCX file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

title = doc.add_heading('Go Language Basics', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph('A concise introduction to the Go programming language (~400 words).', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(64, 64, 64)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)

def body(text):
    return doc.add_paragraph(text, style='Normal')

# ──────────────────────────────────────

add_heading('1. What is Go?', level=1)
body(
    'Go (Golang) is a statically typed, compiled language created at Google by Robert Griesemer, '
    'Rob Pike, and Ken Thompson. Released in 2009, it blends C-like performance with the ease of '
    'dynamic languages. With only 25 keywords, Go is deliberately minimal. Its standout features '
    'include goroutines for lightweight concurrency, channels for communication, a comprehensive '
    'standard library, and built-in tooling for formatting, testing, and dependency management.'
)

add_heading('2. Variables and Types', level=1)
body(
    'Declare variables with var or := (the short operator, usable only inside functions). Go '
    'infers types from initializers. Basic types include bool, string, int (and sized variants), '
    'uint, float32, and float64. All type conversions are explicit. Zero values: 0 for numbers, '
    '"" for strings, false for booleans.'
)
add_code('var name string = "Alice"\nage := 30\ncity := "New York"')

add_heading('3. Control Flow', level=1)
body(
    'Go provides if, else, switch, and for — the only loop construct. Parentheses are omitted '
    'but braces are mandatory. An if statement may include a short initialization before the '
    'condition. Switch cases break automatically and can match multiple values or expressions.'
)
add_code('if x := compute(); x > 10 {\n    fmt.Println("large")\n}\nfor sum < 1000 { sum += sum }')

add_heading('4. Functions', level=1)
body(
    'Functions use the func keyword and support multiple return values — a common pattern is '
    'returning result, error. Named returns allow bare returns. Functions are first-class citizens. '
    'defer schedules a call to run when the enclosing function exits, ideal for cleanup.'
)
add_code('func divide(a, b float64) (float64, error) {\n    if b == 0 { return 0, errors.New("division by zero") }\n    return a / b, nil\n}')

add_heading('5. Structs, Methods & Interfaces', level=1)
body(
    'Go has no classes. Structs hold data; methods attach to any type via a receiver (pointer '
    'receivers for mutation, value receivers for copies). Interfaces are satisfied implicitly — '
    'a type implements an interface by implementing its methods. The error interface underpins '
    'idiomatic error handling.'
)
add_code('type Person struct { Name string; Age int }\nfunc (p *Person) Birthday() { p.Age++ }')

add_heading('6. Concurrency', level=1)
body(
    'Launch a goroutine with go — thousands can run concurrently. Channels enable typed, '
    'safe communication between goroutines (buffered or unbuffered). select multiplexes '
    'channel operations, supporting timeouts and non-blocking patterns.'
)
add_code('ch := make(chan string)\ngo func() { ch <- "hello" }()\nfmt.Println(<-ch)')

add_heading('7. Tooling & Getting Started', level=1)
body(
    'Go\'s toolchain includes go build, run, test, fmt, mod, and more. The standard library '
    'covers HTTP, JSON, crypto, and I/O. To begin: download from go.dev, run go mod init, '
    'write .go files, and use go run. Visit tour.golang.org for an interactive introduction.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('— End of Document —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = '/home/nan/agentos-owl/apps/desktop/Go_Language_Basics.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
